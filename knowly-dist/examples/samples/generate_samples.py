#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
知了 knowly · 测试样例文档生成脚本
=====================================
生成一套自造、零版权的中文测试样例，覆盖知了 v0.1 各项验收能力：
  - 文字版 PDF（单栏 + 双栏，中英混排）—— 测 PDF 解析 + 版面分析（双栏识别）
  - 扫描版 PDF（图片型）                —— 测扫描版检测 + OCR
  - Word（.docx）                       —— 测 Word 解析
  - 图片（.png，含中文）                —— 测 OCR
  - 重复文件（内容相同不同名）          —— 测去重
  - 损坏文件（损坏 PDF）                —— 测失败隔离
  - 繁体版                              —— 测繁简统一 + 近似去重

PDF 一律用 PyMuPDF (fitz) 生成 —— 它原生支持 CJK 字体插入且有真实文本层，
比 fpdf2 / reportlab 在中文字体处理上更稳定。

内容均为自造（《知了示例·茶文化入门》《知了示例·常见算法速查》），无任何真实版权内容。

用法：
    pip install python-docx Pillow PyMuPDF fonttools
    python generate_samples.py
"""

import os
import shutil
import textwrap
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

# ────────────────────────────────────────────────────────────
# 0. 路径与字体
# ────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.resolve()
FONT_DIR = Path("/data/knowly/.fonts")
CN_FONT = str(FONT_DIR / "NotoSansSC-Regular.otf")
CN_FONT_BOLD = str(FONT_DIR / "NotoSansSC-Bold.otf")

# A4 尺寸（PDF point，1pt = 1/72 inch；A4 = 595 x 842 pt）
A4_W, A4_H = 595, 842


def prepare_fonts():
    """若单文件字体不存在，从系统 Noto CJK ttc 抽取 SC 子字体为 otf。"""
    if os.path.exists(CN_FONT) and os.path.exists(CN_FONT_BOLD):
        return
    print("准备中文字体（从 Noto CJK ttc 抽取 SC 子字体）...")
    from fontTools.ttLib import TTCollection
    FONT_DIR.mkdir(parents=True, exist_ok=True)
    for ttc, out_name in [
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansSC-Regular.otf"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc", "NotoSansSC-Bold.otf"),
    ]:
        out_path = FONT_DIR / out_name
        if out_path.exists() or not os.path.exists(ttc):
            continue
        col = TTCollection(ttc)
        target = next((i for i, f in enumerate(col.fonts)
                       if "SC" in (f['name'].getDebugName(1) or "")), 0)
        col.fonts[target].save(str(out_path))
        print(f"  ✓ {out_name}")


# ────────────────────────────────────────────────────────────
# 1. 自造内容（零版权）
# ────────────────────────────────────────────────────────────
TEA_CONTENT = {
    "title": "知了示例·茶文化入门",
    "sections": [
        {
            "h1": "第一章 茶的起源",
            "blocks": [
                ("h2", "一、神农尝百草"),
                ("p", "茶，发乎神农，闻于鲁周公。据《茶经》记载，神农尝百草，日遇七十二毒，得茶而解之。"
                       "这是关于茶最早的传说，也是茶文化在中国的起点。茶由药用渐成饮品，"
                       "历经唐代煮茶、宋代点茶、明代泡茶三大阶段，最终形成今日的饮茶方式。"),
                ("p", "需要说明的是，所谓「神农尝百草」更多是文化象征，并无确凿考古证据。"
                       "但至少可以确定，早在西汉时期，茶已是中国西南地区的日常饮品。"),
                ("h2", "二、茶经与陆羽"),
                ("p", "唐代陆羽（733—804）所著《茶经》三卷，是世界上第一部茶学专著，"
                       "系统记载了茶的起源、采制、烹饮、器具与历史。陆羽因此被后世尊为「茶圣」。"
                       "《茶经》的出现，标志着饮茶从日常生活上升为一门独立的学问。"),
            ],
        },
        {
            "h1": "第二章 六大茶类",
            "blocks": [
                ("p", "中国茶按发酵程度，传统上分为六大类。它们虽然都来自茶树（Camellia sinensis）的叶子，"
                       "却因加工工艺不同，呈现出截然不同的色香味。"),
                ("h2", "一、绿茶（不发酵）"),
                ("p", "绿茶是保留了茶叶天然色泽的茶类，发酵度接近零。杀青是绿茶加工的关键工序，"
                       "通过高温迅速破坏酶的活性，阻止多酚类物质氧化。代表茶有西湖龙井、碧螺春、信阳毛尖。"),
                ("list", ["西湖龙井：产于浙江杭州，扁平挺直，色翠香郁",
                          "碧螺春：产于江苏苏州，条索紧结，卷曲成螺",
                          "信阳毛尖：产于河南信阳，细圆紧直，白毫显露"]),
                ("h2", "二、红茶（全发酵）"),
                ("p", "红茶是全发酵茶，发酵度 80%—95%。在英语世界，「Black Tea」即指红茶。"
                       "发酵过程中，茶多酚氧化生成茶黄素与茶红素，形成红茶特有的红汤红叶。"
                       "代表茶有正山小种、祁门红茶、滇红。"),
                ("list", ["正山小种：世界红茶鼻祖，产于福建武夷山，带松烟香",
                          "祁门红茶：产于安徽祁门，「祁门香」似花似果",
                          "滇红：产于云南，芽叶肥壮，滋味浓厚"]),
                ("h2", "三、乌龙茶（半发酵）"),
                ("p", "乌龙茶又称青茶，属半发酵茶，发酵度介于绿茶与红茶之间，约 10%—70%。"
                       "「做青」是乌龙茶加工的核心工序，通过摇青促使叶缘相互摩擦，"
                       "发生局部的酶促氧化，形成「绿叶红镶边」的特征。代表茶有铁观音、大红袍、凤凰单丛。"),
            ],
        },
        {
            "h1": "第三章 饮茶与健康",
            "blocks": [
                ("p", "现代研究表明，茶叶中含有茶多酚、咖啡碱、茶氨酸、维生素等多种活性成分。"
                       "适量饮茶有助于提神醒脑、抗氧化。但需注意，饮茶也讲究时令与体质，"
                       "空腹饮浓茶、睡前饮茶、服药时饮茶，均可能带来不适。"),
                ("p", "（本节为示例文本，不构成任何医学建议。任何健康问题请咨询专业人士。）"),
            ],
        },
    ],
}

ALGO_CONTENT = {
    "title": "知了示例·常见算法速查",
    "sections": [
        {
            "h1": "Chapter 1 Sorting",
            "blocks": [
                ("p", "Sorting is a fundamental operation in computer science. "
                       "排序是计算机科学中的基础操作。下面介绍两种经典排序算法。"),
                ("h2", "1.1 Quick Sort 快速排序"),
                ("p", "Quick Sort uses divide-and-conquer. Average time complexity is O(n log n), "
                       "worst case O(n²). 快速排序采用分治思想，平均时间复杂度 O(n log n)，最坏 O(n²)。"),
                ("code", "def quicksort(arr):\n"
                         "    if len(arr) <= 1:\n"
                         "        return arr\n"
                         "    pivot = arr[0]\n"
                         "    left = [x for x in arr[1:] if x < pivot]\n"
                         "    right = [x for x in arr[1:] if x >= pivot]\n"
                         "    return quicksort(left) + [pivot] + quicksort(right)"),
                ("h2", "1.2 Merge Sort 归并排序"),
                ("p", "Merge Sort is stable with guaranteed O(n log n) complexity. "
                       "归并排序稳定，时间复杂度稳定为 O(n log n)。空间复杂度 O(n)。"),
            ],
        },
        {
            "h1": "Chapter 2 Searching",
            "blocks": [
                ("h2", "2.1 Binary Search 二分查找"),
                ("p", "Binary Search requires a sorted array. Time complexity O(log n). "
                       "二分查找要求数组有序，时间复杂度 O(log n)。"),
                ("code", "def binary_search(arr, target):\n"
                         "    lo, hi = 0, len(arr) - 1\n"
                         "    while lo <= hi:\n"
                         "        mid = (lo + hi) // 2\n"
                         "        if arr[mid] == target:\n"
                         "            return mid\n"
                         "        elif arr[mid] < target:\n"
                         "            lo = mid + 1\n"
                         "        else:\n"
                         "            hi = mid - 1\n"
                         "    return -1"),
            ],
        },
    ],
}

TRADITIONAL_CONTENT = {
    "title": "知瞭示例·茶文化入門（繁體版）",
    "sections": [
        {
            "h1": "第一章 茶的起源",
            "blocks": [
                ("h2", "一、神農嘗百草"),
                ("p", "茶，發乎神農，聞於魯周公。據《茶經》記載，神農嘗百草，"
                       "日遇七十二毒，得茶而解之。這是關於茶最早的傳說，"
                       "也是茶文化在中國的起點。"),
                ("p", "需要說明的是，所謂「神農嘗百草」更多是文化象徵，並無確鑿考古證據。"
                       "但至少可以確定，早在西漢時期，茶已是中國西南地區的日常飲品。"),
            ],
        },
    ],
}


# ────────────────────────────────────────────────────────────
# 2. PyMuPDF 文本绘制辅助
# ────────────────────────────────────────────────────────────
def font_for(bold):
    return ("NotoSCB", CN_FONT_BOLD) if bold else ("NotoSC", CN_FONT)


def setup_page(page, fonts):
    """注册字体到页面。"""
    for alias, path in fonts:
        page.insert_font(fontname=alias, fontfile=path)


def draw_wrapped_text(page, x, y, text, font_obj, font_alias, fontsize, max_width, line_height):
    """在指定宽度内绘制自动换行的文本，返回结束 y 坐标。font_obj 用于测宽，font_alias/fontfile 用于绘制。"""
    line = ""
    for ch in text:
        test = line + ch
        if font_obj.text_length(test, fontsize=fontsize) > max_width and line:
            page.insert_text((x, y), line, fontname=font_alias, fontsize=fontsize)
            y += line_height
            line = ch
        else:
            line = test
    if line:
        page.insert_text((x, y), line, fontname=font_alias, fontsize=fontsize)
        y += line_height
    return y


def new_doc():
    doc = fitz.open()
    return doc


def render_blocks_single_column(doc, content, margin=50, title_fs=22, h1_fs=16,
                                 h2_fs=13, p_fs=11, list_fs=11, code_fs=9.5,
                                 line_h_factor=1.5):
    """把 content 渲染到 doc（单栏，自动分页）。"""
    # 预加载 Font 对象（用于测宽）
    font_reg = fitz.Font(fontfile=CN_FONT)
    font_bold = fitz.Font(fontfile=CN_FONT_BOLD)

    lh = {title_fs: int(title_fs * line_h_factor), h1_fs: int(h1_fs * line_h_factor),
          h2_fs: int(h2_fs * line_h_factor), p_fs: int(p_fs * line_h_factor),
          list_fs: int(list_fs * line_h_factor), code_fs: int(code_fs * 1.4)}

    def new_page():
        p = doc.new_page(width=A4_W, height=A4_H)
        setup_page(p, [("NotoSC", CN_FONT), ("NotoSCB", CN_FONT_BOLD)])
        return p

    page = new_page()
    max_w = A4_W - 2 * margin
    y = margin + title_fs

    def ensure_space(needed):
        nonlocal page, y
        if y + needed > A4_H - margin:
            page = new_page()
            return margin + p_fs
        return y

    # 标题
    y = ensure_space(lh[title_fs])
    tw = font_bold.text_length(content["title"], fontsize=title_fs)
    page.insert_text(((A4_W - tw) / 2, y), content["title"],
                     fontname="NotoSCB", fontsize=title_fs)
    y += lh[title_fs] + 8

    for section in content["sections"]:
        y = ensure_space(lh[h1_fs] + 10)
        page.insert_text((margin, y), section["h1"], fontname="NotoSCB", fontsize=h1_fs)
        y += lh[h1_fs] + 4
        for kind, text in section["blocks"]:
            if kind == "h2":
                y = ensure_space(lh[h2_fs] + 4)
                page.insert_text((margin, y), text, fontname="NotoSCB", fontsize=h2_fs)
                y += lh[h2_fs] + 2
            elif kind == "p":
                y = ensure_space(lh[p_fs])
                y = draw_wrapped_text(page, margin, y, text, font_reg, "NotoSC",
                                      p_fs, max_w, lh[p_fs])
                y += 4
            elif kind == "list":
                for item in text:
                    y = ensure_space(lh[list_fs])
                    y = draw_wrapped_text(page, margin + 8, y, "· " + item, font_reg, "NotoSC",
                                          list_fs, max_w - 8, lh[list_fs])
                y += 4
            elif kind == "code":
                for line in text.split("\n"):
                    y = ensure_space(lh[code_fs])
                    page.insert_text((margin + 8, y), line if line else " ",
                                     fontname="cour", fontsize=code_fs)
                    y += lh[code_fs]
                y += 4
    return doc


# ────────────────────────────────────────────────────────────
# 3. 单栏文字版 PDF
# ────────────────────────────────────────────────────────────
def make_single_column_pdf(path, content):
    """单栏文字版 PDF，有真实文本层，验证基础 PDF 文本提取。"""
    doc = new_doc()
    render_blocks_single_column(doc, content)
    doc.save(str(path))
    doc.close()
    print(f"  ✓ {path.name}")


# ────────────────────────────────────────────────────────────
# 4. 双栏文字版 PDF（版面分析关键测试样例）
# ────────────────────────────────────────────────────────────
def make_two_column_pdf(path, content):
    """
    双栏文字版 PDF —— 版面分析的关键测试样例。
    文本分置左右两栏区域，验证 RuleBasedLayoutAnalyzer 能否识别双栏
    并按"左栏从上到下 → 右栏从上到下"正确拼接，不串栏。
    含页眉页脚，验证去噪。
    """
    doc = new_doc()
    page = doc.new_page(width=A4_W, height=A4_H)
    setup_page(page, [("NotoSC", CN_FONT), ("NotoSCB", CN_FONT_BOLD)])
    font_reg = fitz.Font(fontfile=CN_FONT)
    font_bold = fitz.Font(fontfile=CN_FONT_BOLD)

    margin = 45
    gutter = 30
    col_w = (A4_W - 2 * margin - gutter) / 2  # 每栏宽
    left_x = margin
    right_x = margin + col_w + gutter

    # 标题横跨两栏
    title_fs = 20
    y = margin + title_fs
    tw = font_bold.text_length(content["title"], fontsize=title_fs)
    page.insert_text(((A4_W - tw) / 2, y), content["title"],
                     fontname="NotoSCB", fontsize=title_fs)
    y += 16
    # 页眉
    header = f"— {content['title']} · 双栏排版示例 —"
    hw = font_reg.text_length(header, fontsize=8)
    page.insert_text(((A4_W - hw) / 2, y),
                     header, fontname="NotoSC", fontsize=8, color=(0.6, 0.6, 0.6))
    col_top = y + 22

    # 把内容平分到左右两栏
    all_blocks = []
    for section in content["sections"]:
        all_blocks.append(("h1", section["h1"]))
        all_blocks.extend(section["blocks"])
    mid = len(all_blocks) // 2

    def draw_column(blocks, x, y_start):
        y = y_start
        h1_fs, h2_fs, p_fs, list_fs = 14, 11, 9.5, 9.5
        lh = lambda fs: int(fs * 1.4)
        for kind, text in blocks:
            if y > A4_H - margin - 20:
                break
            if kind == "h1":
                page.insert_text((x, y), text, fontname="NotoSCB", fontsize=h1_fs)
                y += lh(h1_fs) + 6
            elif kind == "h2":
                page.insert_text((x, y), text, fontname="NotoSCB", fontsize=h2_fs)
                y += lh(h2_fs) + 3
            elif kind == "p":
                y = draw_wrapped_text(page, x, y, text, font_reg, "NotoSC", p_fs, col_w, lh(p_fs))
                y += 3
            elif kind == "list":
                for item in text:
                    y = draw_wrapped_text(page, x + 6, y, "· " + item, font_reg, "NotoSC",
                                          list_fs, col_w - 6, lh(list_fs))
                y += 3
            elif kind == "code":
                for line in text.split("\n"):
                    page.insert_text((x, y), line if line else " ",
                                     fontname="cour", fontsize=8)
                    y += 12
        return y

    draw_column(all_blocks[:mid], left_x, col_top)
    draw_column(all_blocks[mid:], right_x, col_top)

    # 页脚
    footer = f"— 第 1 页 / {content['title']} —"
    fw = font_reg.text_length(footer, fontsize=8)
    page.insert_text(((A4_W - fw) / 2, A4_H - 25),
                     footer, fontname="NotoSC", fontsize=8, color=(0.6, 0.6, 0.6))

    doc.save(str(path))
    doc.close()
    print(f"  ✓ {path.name}（双栏）")


# ────────────────────────────────────────────────────────────
# 5. 扫描版 PDF（图片型，无文本层，测 OCR）
# ────────────────────────────────────────────────────────────
def make_scanned_pdf(path, content):
    """扫描版 PDF：文字渲染成图片放进 PDF，无可提取文本层。"""
    img_path = path.parent / (path.stem + "_src.png")
    render_text_image(img_path, content["title"], content["sections"][0]["blocks"])

    doc = fitz.open()
    page = doc.new_page(width=A4_W, height=A4_H)
    rect = fitz.Rect(40, 50, A4_W - 40, A4_H - 50)
    page.insert_image(rect, filename=str(img_path))
    doc.save(str(path))
    doc.close()
    img_path.unlink()
    print(f"  ✓ {path.name}（扫描版，图片型，无文本层）")


def render_text_image(img_path, title, blocks):
    """把文字渲染成 PNG（模拟扫描件）。"""
    img = Image.new("RGB", (1100, 1500), "white")
    draw = ImageDraw.Draw(img)
    font_title = ImageFont.truetype(CN_FONT, 56)
    font_h2 = ImageFont.truetype(CN_FONT, 36)
    font_p = ImageFont.truetype(CN_FONT, 28)
    y = 60
    draw.text((100, y), title, fill="black", font=font_title)
    y += 90
    for kind, text in blocks[:5]:
        if kind == "h2":
            draw.text((100, y), text, fill="black", font=font_h2)
            y += 50
        elif kind == "p":
            for line in (textwrap.wrap(text, width=28) or [text]):
                draw.text((100, y), line, fill="black", font=font_p)
                y += 38
            y += 20
        elif kind == "list":
            for item in text:
                wrapped = textwrap.wrap(item, width=26) or [item]
                draw.text((100, y), "· " + wrapped[0], fill="black", font=font_p)
                y += 38
                for line in wrapped[1:]:
                    draw.text((130, y), line, fill="black", font=font_p)
                    y += 38
            y += 15
    img.save(str(img_path))


# ────────────────────────────────────────────────────────────
# 6. Word（.docx）
# ────────────────────────────────────────────────────────────
def make_docx(path, content):
    """生成 .docx，验证 Word 解析 + 标题层级保留。"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Noto Sans CJK SC"
    style.font.size = Pt(11)
    doc.add_heading(content["title"], level=0)
    for section in content["sections"]:
        doc.add_heading(section["h1"], level=1)
        for kind, text in section["blocks"]:
            if kind == "h2":
                doc.add_heading(text, level=2)
            elif kind == "p":
                doc.add_paragraph(text)
            elif kind == "list":
                for item in text:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "code":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
    doc.save(str(path))
    print(f"  ✓ {path.name}")


# ────────────────────────────────────────────────────────────
# 7. 含中文的图片（测 OCR）
# ────────────────────────────────────────────────────────────
def make_text_image(path, content):
    """生成一张含中文文字的 PNG，单独测图片 OCR。"""
    img = Image.new("RGB", (1200, 900), "white")
    draw = ImageDraw.Draw(img)
    font_title = ImageFont.truetype(CN_FONT, 48)
    font_p = ImageFont.truetype(CN_FONT, 30)
    y = 50
    draw.text((80, y), content["title"], fill="black", font=font_title)
    y += 80
    test_text = (
        "铁观音产于福建安溪，属乌龙茶类。"
        "正山小种产于福建武夷山，属红茶类，是世界红茶的鼻祖。"
        "龙井产于浙江杭州西湖，属绿茶类，以色翠香郁味甘形美著称。"
        "祁门红茶产于安徽祁门，其特有的「祁门香」似花似果似蜜。"
    )
    for line in textwrap.wrap(test_text, width=30):
        draw.text((80, y), line, fill="black", font=font_p)
        y += 42
    img.save(str(path))
    print(f"  ✓ {path.name}（含中文 OCR 测试）")


# ────────────────────────────────────────────────────────────
# 8. 损坏 PDF（测失败隔离）
# ────────────────────────────────────────────────────────────
def make_corrupted_pdf(path):
    """生成损坏的 PDF（头存在但内容损坏），测失败隔离。"""
    corrupted = ("%PDF-1.4\n%%EOF\n这是一个损坏的 PDF 内容 "
                 "\x00\x01\x02 broken xref table").encode("utf-8", errors="ignore")
    path.write_bytes(corrupted)
    print(f"  ✓ {path.name}（损坏文件，测失败隔离）")


# ────────────────────────────────────────────────────────────
# 9. 主流程
# ────────────────────────────────────────────────────────────
def main():
    prepare_fonts()
    print(f"知了 knowly 测试样例生成（PyMuPDF 方案）")
    print(f"输出目录: {BASE_DIR}")
    print(f"中文字体: {CN_FONT}")
    print()

    print("[主题一：茶文化入门]")
    make_single_column_pdf(BASE_DIR / "tea-single-column.pdf", TEA_CONTENT)
    make_two_column_pdf(BASE_DIR / "tea-two-column.pdf", TEA_CONTENT)
    make_docx(BASE_DIR / "tea.docx", TEA_CONTENT)
    make_scanned_pdf(BASE_DIR / "tea-scanned.pdf", TEA_CONTENT)

    print("\n[主题二：常见算法速查]")
    make_single_column_pdf(BASE_DIR / "algo-single-column.pdf", ALGO_CONTENT)
    make_docx(BASE_DIR / "algo.docx", ALGO_CONTENT)

    print("\n[边界样例]")
    make_text_image(BASE_DIR / "tea-types.png", TEA_CONTENT)
    shutil.copy2(BASE_DIR / "tea.docx", BASE_DIR / "tea-copy-duplicate.docx")
    print(f"  ✓ tea-copy-duplicate.docx（与 tea.docx 内容相同，测文件级去重）")
    make_single_column_pdf(BASE_DIR / "tea-traditional.pdf", TRADITIONAL_CONTENT)
    print(f"  ✓ tea-traditional.pdf（繁体，测繁简统一+近似去重）")
    make_corrupted_pdf(BASE_DIR / "corrupted-broken.pdf")

    print("\n✅ 全部样例生成完成")
    files = [f for f in BASE_DIR.iterdir()
             if f.is_file() and f.name not in ("README.md", "generate_samples.py")]
    print(f"   共 {len(files)} 个样例文件")


if __name__ == "__main__":
    main()
